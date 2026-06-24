import os
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError:
    print("python-docx is not installed yet")
    exit(1)

doc = Document()

# --- JUDUL ---
title = doc.add_heading('Perancangan Database (ERD): Lada Bites', level=0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph('Dokumen ini menjelaskan struktur perancangan database relasional (MySQL) yang digunakan dalam aplikasi Lada Bites. Database ini dirancang dengan prinsip normalisasi untuk mendukung aplikasi E-Commerce skala menengah.')

# --- PENGANTAR RELASI ---
doc.add_heading('1. Konsep Relasi Utama (ERD)', level=1)
doc.add_paragraph('Struktur database Lada Bites memiliki beberapa relasi (hubungan antar tabel) utama:')
doc.add_paragraph('User -> Order (1 to Many): Satu pengguna bisa melakukan banyak pesanan.', style='List Bullet')
doc.add_paragraph('Category -> Product (1 to Many): Satu kategori bisa memiliki banyak produk camilan.', style='List Bullet')
doc.add_paragraph('Order -> OrderItem (1 to Many): Satu pesanan (keranjang) bisa berisi bermacam-macam produk (item).', style='List Bullet')
doc.add_paragraph('Product -> OrderItem (1 to Many): Satu jenis produk bisa dibeli dalam banyak pesanan yang berbeda.', style='List Bullet')
doc.add_paragraph('Product -> Review (1 to Many): Satu produk bisa memiliki banyak ulasan dari pembeli.', style='List Bullet')


# --- DETAIL TABEL ---
doc.add_heading('2. Rincian Tabel Database', level=1)

# Tabel Users
doc.add_heading('Tabel: `users`', level=2)
doc.add_paragraph('Menyimpan data otentikasi dan profil pengguna maupun Admin.')
p = doc.add_paragraph()
p.add_run('Kolom Penting:\n').bold = True
p.add_run('- id (Primary Key)\n- name (Nama Lengkap)\n- email\n- password\n- phone (No. HP)\n- profile_photo_url\n- role (Admin / User biasa)')

# Tabel Categories
doc.add_heading('Tabel: `categories`', level=2)
doc.add_paragraph('Menyimpan daftar kategori produk (contoh: Makaroni, Kerupuk).')
p = doc.add_paragraph()
p.add_run('Kolom Penting:\n').bold = True
p.add_run('- id (Primary Key)\n- name (Nama Kategori)\n- image (Icon Kategori)')

# Tabel Products
doc.add_heading('Tabel: `products`', level=2)
doc.add_paragraph('Menyimpan data katalog produk Lada Bites.')
p = doc.add_paragraph()
p.add_run('Kolom Penting:\n').bold = True
p.add_run('- id (Primary Key)\n- category_id (Foreign Key -> categories.id)\n- nama_produk\n- deskripsi\n- harga\n- stok\n- gambar (URL/Path gambar produk)')

# Tabel Orders
doc.add_heading('Tabel: `orders`', level=2)
doc.add_paragraph('Menyimpan informasi utama sebuah transaksi/pesanan.')
p = doc.add_paragraph()
p.add_run('Kolom Penting:\n').bold = True
p.add_run('- id (Primary Key)\n- user_id (Foreign Key -> users.id)\n- tanggal_pesan\n- total_harga\n- status (Menunggu Pembayaran, Menunggu Konfirmasi, Dikirim, Selesai)\n- metode_pembayaran (QRIS / COD)\n- bukti_pembayaran\n- nama_penerima\n- no_hp\n- alamat_lengkap\n- wilayah_pengiriman\n- ongkir (Tarif pengiriman)')

# Tabel Order Items
doc.add_heading('Tabel: `order_items`', level=2)
doc.add_paragraph('Tabel Pivot/Penghubung yang memecah satu pesanan (Order) menjadi rincian produk apa saja yang dibeli (Keranjang).')
p = doc.add_paragraph()
p.add_run('Kolom Penting:\n').bold = True
p.add_run('- id (Primary Key)\n- order_id (Foreign Key -> orders.id)\n- product_id (Foreign Key -> products.id)\n- jumlah (Quantity)\n- harga_satuan\n- subtotal (jumlah * harga_satuan)')

# Tabel Shipping Rates
doc.add_heading('Tabel: `shipping_rates`', level=2)
doc.add_paragraph('Menyimpan daftar tarif ongkos kirim berdasarkan wilayah/daerah.')
p = doc.add_paragraph()
p.add_run('Kolom Penting:\n').bold = True
p.add_run('- id (Primary Key)\n- wilayah (Nama Daerah)\n- tarif (Harga Ongkir)')

# Tabel Reviews
doc.add_heading('Tabel: `reviews`', level=2)
doc.add_paragraph('Menyimpan ulasan/bintang dari pelanggan yang telah membeli produk.')
p = doc.add_paragraph()
p.add_run('Kolom Penting:\n').bold = True
p.add_run('- id (Primary Key)\n- user_id (Foreign Key -> users.id)\n- product_id (Foreign Key -> products.id)\n- rating (Bintang 1-5)\n- comment (Komentar ulasan)')

# Tabel Banners
doc.add_heading('Tabel: `banners`', level=2)
doc.add_paragraph('Menyimpan data gambar promosi yang akan bergeser (slider) di halaman utama aplikasi (Home Screen).')
p = doc.add_paragraph()
p.add_run('Kolom Penting:\n').bold = True
p.add_run('- id (Primary Key)\n- title (Judul Promo)\n- image_url\n- is_active (Status aktif/tidak)')


# --- KESIMPULAN ---
doc.add_heading('3. Alur Logika Database Saat Checkout', level=1)
doc.add_paragraph('1. Saat pelanggan menekan tombol Checkout, sistem akan membuat 1 baris data di tabel `orders` yang mencatat informasi pengiriman dan total biaya.', style='List Number')
doc.add_paragraph('2. Secara bersamaan, sistem akan melakukan perulangan (looping) untuk menyimpan setiap produk di keranjang ke dalam tabel `order_items`, dengan mengaitkan `order_id` dari langkah pertama.', style='List Number')
doc.add_paragraph('3. Saat pengguna mengirim bukti WA, admin memverifikasi dan sistem hanya perlu memperbarui kolom `status` pada tabel `orders` tanpa menyentuh tabel `order_items`.', style='List Number')


# Simpan Dokumen
file_path = "Perancangan_Database_Lada_Bites.docx"
doc.save(file_path)
print(f"File berhasil dibuat di: {os.path.abspath(file_path)}")
