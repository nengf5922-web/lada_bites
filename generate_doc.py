import os
try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError:
    print("python-docx is not installed yet")
    exit(1)

doc = Document()

# Judul
title = doc.add_heading('Dokumentasi Fitur Aplikasi Lada Bites', level=0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph('Dokumen ini berisi rangkuman bagian dan fitur penting dari setiap file utama dalam aplikasi Lada Bites (Frontend & Backend).')

doc.add_heading('1. Struktur Utama Aplikasi', level=1)
doc.add_paragraph('Aplikasi Lada Bites terbagi menjadi dua bagian utama:')
doc.add_paragraph('- Frontend: Dibangun menggunakan Flutter, berfungsi sebagai antarmuka pengguna (pembeli) di aplikasi mobile.', style='List Bullet')
doc.add_paragraph('- Backend: Dibangun menggunakan Laravel (PHP), bertindak sebagai penyedia API (Application Programming Interface) serta panel Admin untuk mengelola data.', style='List Bullet')

doc.add_heading('2. Frontend (Flutter) - Fitur Utama', level=1)

# Home Screen
doc.add_heading('Home Screen (home_screen.dart)', level=2)
doc.add_paragraph('Fungsi: Halaman beranda utama untuk melihat katalog produk dan promosi.')
p = doc.add_paragraph('Fitur Penting:')
doc.add_paragraph('Carousel Banner: Menampilkan promo/diskon secara dinamis.', style='List Bullet 2')
doc.add_paragraph('Grid Produk: Mengambil data produk langsung dari Backend API (ProductApiService) dan menampilkannya dengan foto, nama, dan harga.', style='List Bullet 2')
doc.add_paragraph('Pencarian: Membantu pembeli mencari cemilan berdasarkan nama secara real-time.', style='List Bullet 2')

# Keranjang
doc.add_heading('Keranjang Belanja (cart_screen.dart & cart_provider.dart)', level=2)
doc.add_paragraph('Fungsi: Tempat menampung produk sebelum dilakukan proses checkout.')
p = doc.add_paragraph('Fitur Penting:')
doc.add_paragraph('State Management (Provider): Digunakan agar jumlah item keranjang selalu up-to-date di seluruh halaman secara otomatis.', style='List Bullet 2')
doc.add_paragraph('Logika Checkbox Dinamis: Produk yang baru ditambahkan ke keranjang secara otomatis tidak dicentang (harus dipilih manual) agar pembeli tidak keliru saat checkout.', style='List Bullet 2')

# Checkout
doc.add_heading('Checkout Screen (checkout_screen.dart)', level=2)
doc.add_paragraph('Fungsi: Halaman konfirmasi pemesanan, pengisian alamat, dan pemilihan metode pembayaran.')
p = doc.add_paragraph('Fitur Penting:')
doc.add_paragraph('Perhitungan Total Cerdas: Menghitung otomatis total harga produk ditambah ongkos kirim berdasarkan wilayah yang dipilih.', style='List Bullet 2')
doc.add_paragraph('Data Alamat: Mengambil alamat profil pengguna secara langsung atau memungkinkan perubahan alamat sementara khusus untuk pesanan tersebut.', style='List Bullet 2')
doc.add_paragraph('Kirim Pesanan (API): Data keranjang yang valid akan dikirimkan ke Laravel untuk dicatat ke dalam database secara real-time.', style='List Bullet 2')

# Payment
doc.add_heading('Payment Screen (payment_screen.dart)', level=2)
doc.add_paragraph('Fungsi: Halaman instruksi pembayaran dan proses konfirmasi bukti transfer.')
p = doc.add_paragraph('Fitur Penting:')
doc.add_paragraph('QRIS Display: Menampilkan barcode pembayaran QRIS resmi toko secara jelas lengkap dengan peringatan anti-penipuan.', style='List Bullet 2')
doc.add_paragraph('Otomatisasi WhatsApp: Menggantikan fitur unggah gambar manual yang merepotkan menjadi satu tombol praktis yang otomatis melompat membuka aplikasi WhatsApp lengkap dengan format template pesanan.', style='List Bullet 2')
doc.add_paragraph('Auto-Status Update: Sesaat setelah tombol WhatsApp ditekan, aplikasi akan secara background memanggil API untuk mengubah status pesanan dari "Menunggu Pembayaran" menjadi "Menunggu Konfirmasi".', style='List Bullet 2')

# History
doc.add_heading('History Screen (history_screen.dart)', level=2)
doc.add_paragraph('Fungsi: Melihat daftar seluruh riwayat pesanan yang telah dibuat oleh pembeli.')
p = doc.add_paragraph('Fitur Penting:')
doc.add_paragraph('Tracking Status: Menampilkan label status terkini (Belum Dibayar, Menunggu Konfirmasi, Dikirim, Selesai).', style='List Bullet 2')
doc.add_paragraph('Smart Image Parser: Memiliki kecerdasan untuk menangani kegagalan pemuatan gambar; aplikasi bisa membedakan mana gambar yang berasal dari aset lokal (assets/...) dan mana yang dari jaringan/server backend.', style='List Bullet 2')

doc.add_heading('3. Backend (Laravel) - Fitur Utama', level=1)

# Order Controller
doc.add_heading('Order Controller (OrderController.php)', level=2)
doc.add_paragraph('Fungsi: Otak utama di server dari keseluruhan proses transaksi.')
p = doc.add_paragraph('Fitur Penting:')
doc.add_paragraph('Pemrosesan Transaksi: Menerima keranjang belanja dari Flutter dan secara aman memasukkannya ke tabel `orders` dan `order_items` di MySQL.', style='List Bullet 2')
doc.add_paragraph('Endpoint Status: Menyediakan jalur komunikasi khusus untuk aplikasi mengubah status pesanan tanpa celah keamanan.', style='List Bullet 2')

# Admin Panel
doc.add_heading('Admin Panel (Blade Views - pesanan.blade.php)', level=2)
doc.add_paragraph('Fungsi: Antarmuka bagi pemilik/admin toko untuk mengelola dan mengawasi bisnis.')
p = doc.add_paragraph('Fitur Penting:')
doc.add_paragraph('Dashboard Pesanan: Halaman khusus tempat Admin melihat daftar pesanan masuk secara real-time, lalu memverifikasinya setelah menerima bukti pembayaran di WhatsApp.', style='List Bullet 2')

# Simpan Dokumen
file_path = "Dokumentasi_Aplikasi_Lada_Bites.docx"
doc.save(file_path)
print(f"File berhasil dibuat di: {os.path.abspath(file_path)}")
