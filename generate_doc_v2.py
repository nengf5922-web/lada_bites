import os
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError:
    print("python-docx is not installed yet")
    exit(1)

doc = Document()

# --- JUDUL ---
title = doc.add_heading('Buku Panduan & Dokumentasi Sistem: Lada Bites', level=0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph('Dokumen ini berisi penjelasan komprehensif mengenai struktur file, arsitektur REST API, serta daftar pertanyaan teknis yang sering ditanyakan terkait sistem aplikasi Lada Bites.')

# --- BAGIAN 1 ---
doc.add_heading('BAB 1: Arsitektur REST API & Format JSON', level=1)
doc.add_paragraph('Aplikasi Lada Bites menggunakan arsitektur REST API (Representational State Transfer) untuk menghubungkan aplikasi Mobile (Flutter) dengan Server (Laravel).')

doc.add_heading('Mengapa menggunakan format JSON?', level=2)
p = doc.add_paragraph()
p.add_run('JSON (JavaScript Object Notation)').bold = True
p.add_run(' dipilih sebagai format pertukaran data utama dengan alasan berikut:')
doc.add_paragraph('1. Ringan dan Cepat: Sintaks JSON sangat minim karakter dibandingkan XML, sehingga ukuran payload/data yang dikirim lewat internet menjadi sangat kecil. Hal ini menghemat kuota pengguna dan mempercepat proses loading aplikasi.', style='List Bullet')
doc.add_paragraph('2. Mudah Di-parsing: Bahasa Dart (Flutter) memiliki dukungan bawaan yang sangat baik untuk menerjemahkan objek JSON menjadi Model Dart.', style='List Bullet')
doc.add_paragraph('3. Standar Industri: JSON adalah standar tak tertulis di industri pengembangan perangkat lunak modern untuk komunikasi REST API.', style='List Bullet')


# --- BAGIAN 2 ---
doc.add_heading('BAB 2: Penjelasan Struktur File', level=1)

# Frontend
doc.add_heading('A. Frontend (Flutter) - Folder lib/', level=2)
doc.add_paragraph('main.dart', style='List Bullet').bold = True
doc.add_paragraph('File utama/titik masuk aplikasi. Bertugas menginisialisasi rute, Provider (state management), dan tema global aplikasi.')
doc.add_paragraph('screens/home_screen.dart', style='List Bullet').bold = True
doc.add_paragraph('Halaman beranda aplikasi. Mengambil data produk dari API, menampilkan carousel banner promo, serta daftar kategori.')
doc.add_paragraph('screens/cart_screen.dart & providers/cart_provider.dart', style='List Bullet').bold = True
doc.add_paragraph('Halaman keranjang belanja. Menggunakan "Provider" agar jumlah barang bisa sinkron di seluruh aplikasi tanpa harus me-refresh halaman (Reactive State).')
doc.add_paragraph('screens/checkout_screen.dart', style='List Bullet').bold = True
doc.add_paragraph('Halaman penyelesaian pesanan. Melakukan perhitungan ongkos kirim otomatis dan merangkum payload JSON sebelum dikirim ke Laravel.')
doc.add_paragraph('screens/payment_screen.dart', style='List Bullet').bold = True
doc.add_paragraph('Halaman instruksi QRIS. Dilengkapi integrasi `url_launcher` untuk melompat secara otomatis ke aplikasi WhatsApp Admin, sekaligus mengupdate status pesanan secara background.')
doc.add_paragraph('screens/history_screen.dart', style='List Bullet').bold = True
doc.add_paragraph('Menampilkan riwayat belanja pengguna. Menggunakan Smart Image Parser untuk mendeteksi apakah gambar harus diload dari memori HP (assets) atau dari internet (API).')
doc.add_paragraph('services/ (contoh: order_api_service.dart)', style='List Bullet').bold = True
doc.add_paragraph('Kumpulan file kurir. Menggunakan package "Dio" untuk membungkus data menjadi HTTP Request (GET, POST) dan mengirimkannya ke server backend.')

# Backend
doc.add_heading('B. Backend (Laravel) - Server', level=2)
doc.add_paragraph('routes/api.php & routes/web.php', style='List Bullet').bold = True
doc.add_paragraph('Pusat terminal lalu lintas. api.php mengatur rute untuk aplikasi Flutter (menggunakan proteksi token Sanctum), sedangkan web.php untuk rute website Admin.')
doc.add_paragraph('app/Http/Controllers/OrderController.php', style='List Bullet').bold = True
doc.add_paragraph('Otak logika transaksi. Menerima JSON dari Flutter, memvalidasi stok, lalu menyimpannya ke tabel `orders` dan `order_items` di MySQL.')
doc.add_paragraph('app/Models/ (contoh: Order.php, Product.php)', style='List Bullet').bold = True
doc.add_paragraph('Representasi (ORM) dari tabel-tabel di database MySQL. Membantu backend memanggil data tanpa harus menulis query SQL mentah.')
doc.add_paragraph('resources/views/admin/pesanan.blade.php', style='List Bullet').bold = True
doc.add_paragraph('Halaman antarmuka khusus pemilik toko (Admin) yang dibangun dengan HTML/Blade untuk memantau, menerima, atau menolak pesanan pembeli.')


# --- BAGIAN 3 ---
doc.add_heading('BAB 3: Kemungkinan Pertanyaan Presentasi / Sidang (Q&A)', level=1)

qna_list = [
    ("Q: Mengapa menggunakan Flutter untuk mobile dan Laravel untuk Backend?",
     "A: Flutter memungkinkan pembuatan aplikasi Android dan iOS sekaligus dari satu kode (cross-platform), sehingga sangat efisien dari segi waktu dan biaya. Laravel dipilih karena memiliki ekosistem bawaan yang matang (seperti Eloquent ORM, migrasi database, dan proteksi API Sanctum) untuk membangun server REST API yang stabil dengan cepat."),
    
    ("Q: Bagaimana alur pengamanan data pengguna dalam sistem ini?",
     "A: Aplikasi menggunakan sistem Token-Based Authentication dengan Laravel Sanctum. Saat pengguna berhasil login, server memberikan kunci unik (Token). Token ini selalu diselipkan di setiap request (pada Header HTTP) oleh package Dio di Flutter. Tanpa token ini, server akan memblokir request tersebut."),
    
    ("Q: Mengapa verifikasi pembayaran dialihkan ke WhatsApp daripada membuat sistem upload foto struk di dalam aplikasi?",
     "A: Ini adalah strategi penyederhanaan User Experience (UX) dan efisiensi Server. Sistem upload gambar dari galeri HP seringkali rentan terhadap bug kompatibilitas versi Android, dan memakan kapasitas storage server yang besar. Dengan WhatsApp, pelanggan merasa lebih dekat secara psikologis dengan CS, prosesnya terjamin berhasil, dan server tidak terbebani file sampah."),
     
    ("Q: Jika aplikasi Flutter ditutup secara paksa, apakah data keranjang belanja akan hilang?",
     "A: Untuk arsitektur saat ini yang menggunakan Provider in-memory, data keranjang bersifat sementara. Hal ini bisa ditingkatkan di masa depan dengan menggunakan penyimpanan lokal seperti SharedPreferences atau SQLite agar keranjang persisten (tersimpan permanen)."),
     
    ("Q: Apa tantangan terbesar saat menghubungkan Flutter ke Laravel?",
     "A: Salah satu tantangannya adalah penanganan URL Localhost. Emulator Android (Flutter) tidak bisa membaca URL '127.0.0.1' karena mengacu pada perangkat emulator itu sendiri. Hal ini diatasi dengan me-replace URL gambar secara dinamis menjadi '10.0.2.2' khusus ketika mendeteksi lingkungan Android.")
]

for q, a in qna_list:
    pq = doc.add_paragraph()
    pq.add_run(q).bold = True
    pq.add_run('\n' + a)

# Simpan Dokumen
file_path = "Dokumentasi_Lengkap_Lada_Bites.docx"
doc.save(file_path)
print(f"File berhasil dibuat di: {os.path.abspath(file_path)}")
